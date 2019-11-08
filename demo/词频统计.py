import docx
import requests
import urllib.request
from lxml import etree
import re
import time
from functools import reduce

doc_file_path1 = r'C:\Users\Administrator\Desktop\单词总览.docx'
doc_file_path2 = r'C:\Users\Administrator\Desktop\词汇总结.docx'

doc_file1 = docx.Document(doc_file_path1)
doc_file2 = docx.Document(doc_file_path2)

word_list = [] # 统计单词
# word_set1 = set() # 统计需注释的单词
# word_set2 = set() # 统计已注释的单词

# word_add 函数（文件,统计单词集合名称，关键字段）
def word_add (doc_file,key_str=''):
    word_set = set()
    for para in doc_file.paragraphs:
        # print(para.text)
        for word in para.text.split(' '):
            # print(word)
            if key_str !='' and key_str in word :
                word_set.add(word.split(key_str)[0])
            else:
                word_set.add(word)
    return word_set
# print(word_set1)

def word_translate(word):
    data = {
    'doctype': 'json',
    'type': 'AUTO',
    'i':word
    }
    url = "http://fanyi.youdao.com/translate"
    r = requests.get(url,params=data)
    result = r.json()
    # print(result)
    translate_result = result['translateResult'][0][0]["tgt"]
    # print(translate_result,type(translate_result))
    return translate_result

#获得页面数据
def get_page(myword):
    basurl = "http://cn.bing.com/dict/search?q="
    searchurl=basurl+myword
    response =  urllib.request.urlopen(searchurl)
    html = response.read()
    return html

#获得单词释义
def get_chitiao(html_selector):
    chitiao=[]
    hanyi_xpath='/html/body/div[1]/div/div/div[1]/div[1]/ul/li'
    get_hanyi=html_selector.xpath(hanyi_xpath)
    for item in get_hanyi:
        it=item.xpath('span')
        chitiao.append('%s||%s'%(it[0].text,it[1].xpath('span')[0].text))
    if len(chitiao)>0:
        # print(chitiao)
        return reduce(lambda x, y:"%s||||%s"%(x,y),chitiao)
    else:
        return ""


word_set1 = word_add(doc_file1)
word_set2 = word_add(doc_file2)
word_set = word_set1.difference(word_set2)
# print (word_set1)
# print (word_set2)
# print (word_set)
# print (word_translate('word'))

def word_translate_add(word_set,file_save):
    print(word_set)
    for word in word_set:
        #获得页面
        pagehtml=get_page(word)
        selector = etree.HTML(pagehtml.decode('utf-8'))
        #单词释义
        chitiao=get_chitiao(selector)
        file = open(file_save,'w')

        text = (lambda chitiao : word+' '+chitiao if chitiao !='' else word+' 无释义，注意拼写'),(chitiao)
        # if chitiao == '':
        #     text = print(word,'无释义，注意拼写')
        # else:
        #     text = print(word,chitiao)
        file.writelines(word+' '+text[1])
        # print(word+''+text[1])
        print (word)
        file.close()


try:
    word_translate_add(word_set,r'C:\Users\Administrator\Desktop\2.txt')
except Exception as e:
    print('Error:',e)
finally:
    print('报错了，看一下')